"""
SAARJ Word → LaTeX Formatter
Converts free-format Word documents to SAARJ journal LaTeX template.
Also supports structured form-based input (generate_latex_from_form).
"""

import re
import os
import io
import zipfile
from docx import Document
from docx.oxml.ns import qn


# ── LaTeX special-character escaping ──────────────────────────────────────────
_LATEX_ESCAPE = str.maketrans({
    '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#',
    '_': r'\_', '{': r'\{', '}': r'\}',
    '~': r'\textasciitilde{}', '^': r'\textasciicircum{}',
    '\\': r'\textbackslash{}',
})

def escape(text: str) -> str:
    if not text:
        return ''
    return text.translate(_LATEX_ESCAPE)


_FORMULA_PATTERN = re.compile(
    r'(\\begin\{(?:equation\*?|align\*?|gather\*?|multline\*?)\}[\s\S]*?\\end\{(?:equation\*?|align\*?|gather\*?|multline\*?)\}'
    r'|\\\[[\s\S]*?\\\]'
    r'|\\\([\s\S]*?\\\)'
    r'|\$\$[\s\S]*?\$\$'
    r'|\$(?:\\.|[^$\\])+\$'
    r'|<math[\s\S]*?</math>)',
    re.MULTILINE
)


def render_rich_text(text: str) -> str:
    """Escape prose while preserving math/formula fragments."""
    if not text:
        return ''

    parts = []
    for chunk in _FORMULA_PATTERN.split(text):
        if not chunk:
            continue
        if _FORMULA_PATTERN.fullmatch(chunk):
            if chunk.lstrip().startswith('<math'):
                parts.append(
                    r'\begin{verbatim}' + '\n' +
                    chunk + '\n' +
                    r'\end{verbatim}'
                )
            else:
                parts.append(chunk)
        else:
            parts.append(escape(chunk))
    return ''.join(parts)


def _local_name(tag: str) -> str:
    return tag.rsplit('}', 1)[-1]


def _omml_to_latex(node) -> str:
    lname = _local_name(node.tag)
    children = list(node)

    if lname in ('oMath', 'oMathPara', 'e', 'num', 'den', 'deg'):
        return ''.join(_omml_to_latex(child) for child in children)
    if lname in ('r', 'mr'):
        return ''.join(_omml_to_latex(child) for child in children)
    if lname == 't':
        return escape(node.text or '')
    if lname == 'sSup':
        base = next((child for child in children if _local_name(child.tag) == 'e'), None)
        sup = next((child for child in children if _local_name(child.tag) == 'sup'), None)
        return '{' + _omml_to_latex(base) + '}^{' + _omml_to_latex(sup) + '}'
    if lname == 'sSub':
        base = next((child for child in children if _local_name(child.tag) == 'e'), None)
        sub = next((child for child in children if _local_name(child.tag) == 'sub'), None)
        return '{' + _omml_to_latex(base) + '}_{' + _omml_to_latex(sub) + '}'
    if lname == 'f':
        num = next((child for child in children if _local_name(child.tag) == 'num'), None)
        den = next((child for child in children if _local_name(child.tag) == 'den'), None)
        return r'\frac{' + _omml_to_latex(num) + '}{' + _omml_to_latex(den) + '}'
    if lname == 'rad':
        deg = next((child for child in children if _local_name(child.tag) == 'deg'), None)
        expr = next((child for child in children if _local_name(child.tag) == 'e'), None)
        if deg is not None:
            return r'\sqrt[' + _omml_to_latex(deg) + ']{' + _omml_to_latex(expr) + '}'
        return r'\sqrt{' + _omml_to_latex(expr) + '}'
    if lname == 'd':
        return '(' + ''.join(_omml_to_latex(child) for child in children) + ')'
    return ''.join(_omml_to_latex(child) for child in children)


# ── Heading / section detection ───────────────────────────────────────────────
SECTION_MAP = {
    # TR keys
    'giriş': 'Giriş / Introduction',
    'introduction': 'Giriş / Introduction',
    'literatür': 'Literatür Taraması / Literature Review',
    'literature': 'Literatür Taraması / Literature Review',
    'literature review': 'Literatür Taraması / Literature Review',
    'kuramsal çerçeve': 'Literatür Taraması / Literature Review',
    'yöntem': 'Yöntem / Methodology',
    'yöntem ve teknik': 'Yöntem / Methodology',
    'methodology': 'Yöntem / Methodology',
    'method': 'Yöntem / Methodology',
    'materials and methods': 'Yöntem / Methodology',
    'bulgular': 'Bulgular / Findings',
    'findings': 'Bulgular / Findings',
    'results': 'Bulgular / Findings',
    'tartışma': 'Tartışma / Discussion',
    'discussion': 'Tartışma / Discussion',
    'sonuç': 'Sonuç / Conclusion',
    'conclusion': 'Sonuç / Conclusion',
    'sonuç ve öneriler': 'Sonuç / Conclusion',
    'conclusions': 'Sonuç / Conclusion',
    'kaynakça': '__REFERENCES__',
    'kaynaklar': '__REFERENCES__',
    'references': '__REFERENCES__',
    'bibliography': '__REFERENCES__',
}

ABSTRACT_KEYS = {'abstract', 'özet', 'öz'}
KEYWORD_KEYS  = {'keywords', 'anahtar kelimeler', 'anahtar sözcükler', 'key words'}


def _para_text(para) -> str:
    return para.text.strip()


def _is_heading(para) -> bool:
    return para.style.name.startswith('Heading') or para.style.name.startswith('Başlık')


def _heading_level(para) -> int:
    name = para.style.name
    for part in name.split():
        if part.isdigit():
            return int(part)
    return 1


def _has_bold(para) -> bool:
    return any(run.bold for run in para.runs if run.text.strip())


def _para_to_latex(para) -> str:
    """Convert a paragraph with inline formatting to LaTeX."""
    has_math = any(_local_name(child.tag) in ('oMath', 'oMathPara', 'object', 'OLEObject') for child in para._p)
    if has_math:
        parts = []
        run_index = 0
        runs = para.runs
        for child in para._p:
            lname = _local_name(child.tag)
            if lname == 'r':
                run = runs[run_index] if run_index < len(runs) else None
                run_index += 1
                if not run:
                    continue
                t = escape(run.text)
                if not t:
                    continue
                if run.bold and run.italic:
                    t = r'\textbf{\textit{' + t + '}}'
                elif run.bold:
                    t = r'\textbf{' + t + '}'
                elif run.italic:
                    t = r'\textit{' + t + '}'
                parts.append(t)
            elif lname in ('oMath', 'oMathPara'):
                math_latex = _omml_to_latex(child).strip()
                if math_latex:
                    parts.append(r'\[' + math_latex + r'\]')
            elif lname in ('object', 'OLEObject'):
                parts.append(r'\begin{equation*}\text{[Embedded equation object preserved from source]}\end{equation*}')
        return ''.join(parts)

    parts = []
    for run in para.runs:
        t = escape(run.text)
        if not t:
            continue
        if run.bold and run.italic:
            t = r'\textbf{\textit{' + t + '}}'
        elif run.bold:
            t = r'\textbf{' + t + '}'
        elif run.italic:
            t = r'\textit{' + t + '}'
        parts.append(t)
    return ''.join(parts)


def _list_to_latex(paras, numbered: bool) -> str:
    env = 'enumerate' if numbered else 'itemize'
    opt = r'[leftmargin=1.2cm, label=\arabic*.]' if numbered else ''
    lines = [r'\begin{' + env + '}' + opt]
    for p in paras:
        lines.append(r'  \item ' + _para_to_latex(p))
    lines.append(r'\end{' + env + '}')
    return '\n'.join(lines)


# ── Word document parser ───────────────────────────────────────────────────────
def extract_from_docx(file_bytes: bytes) -> dict:
    """
    Returns a dict with keys:
      tr_title, en_title, tr_abstract, en_abstract,
      tr_keywords, en_keywords, sections (list of {title, level, latex}),
      references (list of str), raw_paragraphs
    """
    doc = Document(io.BytesIO(file_bytes))
    result = {
        'tr_title': '', 'en_title': '',
        'tr_abstract': '', 'en_abstract': '',
        'tr_keywords': '', 'en_keywords': '',
        'sections': [],
        'references': [],
    }

    paras = [p for p in doc.paragraphs]
    i = 0
    n = len(paras)

    current_section_title = None
    current_section_level = 1
    current_section_lines = []
    in_abstract = False
    in_keywords = False
    abstract_lang = None   # 'tr' or 'en'
    in_references = False

    def flush_section():
        if current_section_title is None:
            return
        result['sections'].append({
            'title': current_section_title,
            'level': current_section_level,
            'latex': '\n\n'.join(current_section_lines),
        })

    # ── first pass: find titles (usually first 1-3 paragraphs before abstract) ──
    first_title_found = False
    for p in paras[:8]:
        txt = _para_text(p)
        if not txt:
            continue
        key = txt.lower().strip().rstrip(':').strip()
        if key in ABSTRACT_KEYS or key in KEYWORD_KEYS:
            break
        if _is_heading(p) or len(txt) > 10:
            if not first_title_found:
                # Heuristic: if line looks like a Turkish title (no ASCII section name)
                if not any(key in txt.lower() for key in SECTION_MAP):
                    result['tr_title'] = escape(txt)
                    first_title_found = True
                    continue
            elif not result['en_title']:
                if not any(key in txt.lower() for key in SECTION_MAP):
                    result['en_title'] = escape(txt)
                    break

    # ── main pass ──
    i = 0
    while i < n:
        para = paras[i]
        txt = _para_text(para)
        key = txt.lower().strip().rstrip(':').strip()
        style = para.style.name

        # Skip empty
        if not txt:
            i += 1
            continue

        # ── References section ──
        if in_references:
            if txt:
                result['references'].append(escape(txt))
            i += 1
            continue

        # ── Heading detection ──
        if _is_heading(para) or (len(txt) < 80 and txt.isupper() and len(txt) > 3):
            mapped = SECTION_MAP.get(key)
            if mapped == '__REFERENCES__':
                flush_section()
                current_section_title = None
                current_section_lines = []
                in_references = True
                i += 1
                continue
            if mapped:
                flush_section()
                current_section_title = mapped
                current_section_level = _heading_level(para) if _is_heading(para) else 1
                current_section_lines = []
                in_abstract = False
                i += 1
                continue
            # Sub-heading inside a known section
            if current_section_title:
                flush_section()
                current_section_title = escape(txt)
                current_section_level = _heading_level(para) if _is_heading(para) else 2
                current_section_lines = []
                i += 1
                continue

        # ── Abstract heading ──
        if key in ABSTRACT_KEYS:
            in_abstract = True
            in_keywords = False
            # determine language from context
            if 'en' in key or key == 'abstract':
                abstract_lang = 'en'
            else:
                abstract_lang = 'tr'
            i += 1
            continue

        # ── Keywords line ──
        if key in KEYWORD_KEYS or txt.lower().startswith('keyword') or txt.lower().startswith('anahtar'):
            in_abstract = False
            in_keywords = True
            # might be on same line: "Keywords: foo, bar"
            colon_pos = txt.find(':')
            if colon_pos != -1:
                kw_val = txt[colon_pos+1:].strip()
                if abstract_lang == 'en':
                    result['en_keywords'] = escape(kw_val)
                else:
                    result['tr_keywords'] = escape(kw_val)
                in_keywords = False
            i += 1
            continue

        if in_keywords:
            if abstract_lang == 'en':
                result['en_keywords'] = escape(txt)
            else:
                result['tr_keywords'] = escape(txt)
            in_keywords = False
            i += 1
            continue

        # ── Abstract body ──
        if in_abstract:
            if abstract_lang == 'en':
                result['en_abstract'] += (' ' if result['en_abstract'] else '') + txt
            else:
                result['tr_abstract'] += (' ' if result['tr_abstract'] else '') + txt
            i += 1
            continue

        # ── Regular paragraph / list ──
        if current_section_title is None:
            i += 1
            continue

        # List paragraph
        if style.startswith('List') or para.style.name in ('List Paragraph', 'Liste Paragrafı'):
            # collect consecutive list items
            list_paras = [para]
            numbered = 'Number' in style or 'Numara' in style
            j = i + 1
            while j < n:
                np2 = paras[j]
                s2 = np2.style.name
                if s2.startswith('List') or s2 in ('List Paragraph',):
                    list_paras.append(np2)
                    j += 1
                else:
                    break
            current_section_lines.append(_list_to_latex(list_paras, numbered))
            i = j
            continue

        # Normal paragraph
        latex_line = _para_to_latex(para)
        if latex_line:
            current_section_lines.append(latex_line)
        i += 1

    flush_section()
    return result


# ── Author info parser ─────────────────────────────────────────────────────────
def parse_author_info(file_bytes: bytes, filename: str) -> list:
    """
    Parse author info from a .docx or .txt file.
    Expected format (one author per line/paragraph):
      Ad Soyad | Kurum | ORCID | email | sorumlu(evet/hayır)
    or Word table with columns: Ad Soyad, Kurum, ORCID, E-posta, Sorumlu
    Returns list of dicts.
    """
    authors = []

    if filename.lower().endswith('.docx'):
        doc = Document(io.BytesIO(file_bytes))
        # Try table first
        if doc.tables:
            tbl = doc.tables[0]
            for row in tbl.rows[1:]:  # skip header
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 4 and cells[0]:
                    authors.append({
                        'name':        cells[0],
                        'affiliation': cells[1] if len(cells) > 1 else '',
                        'orcid':       cells[2] if len(cells) > 2 else '',
                        'email':       cells[3] if len(cells) > 3 else '',
                        'corresponding': len(cells) > 4 and cells[4].lower() in ('evet', 'yes', 'e', 'y', '1', 'true'),
                    })
        else:
            for para in doc.paragraphs:
                txt = para.text.strip()
                if not txt or txt.startswith('#'):
                    continue
                parts = [p.strip() for p in re.split(r'[|\t;]', txt)]
                if len(parts) >= 2:
                    authors.append({
                        'name':        parts[0],
                        'affiliation': parts[1] if len(parts) > 1 else '',
                        'orcid':       parts[2] if len(parts) > 2 else '',
                        'email':       parts[3] if len(parts) > 3 else '',
                        'corresponding': len(parts) > 4 and parts[4].lower() in ('evet', 'yes', 'e', 'y', '1', 'true'),
                    })
    else:
        # Plain text
        text = file_bytes.decode('utf-8', errors='ignore')
        for line in text.splitlines():
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            parts = [p.strip() for p in re.split(r'[|\t;]', line)]
            if len(parts) >= 2:
                authors.append({
                    'name':        parts[0],
                    'affiliation': parts[1] if len(parts) > 1 else '',
                    'orcid':       parts[2] if len(parts) > 2 else '',
                    'email':       parts[3] if len(parts) > 3 else '',
                    'corresponding': len(parts) > 4 and parts[4].lower() in ('evet', 'yes', 'e', 'y', '1', 'true'),
                })

    return authors


# ── LaTeX generation ───────────────────────────────────────────────────────────
SECTION_CMD = {1: r'\section', 2: r'\subsection', 3: r'\subsubsection'}

def _format_author_block(authors: list, corr_marker: str = '*') -> str:
    parts = []
    for idx, a in enumerate(authors, 1):
        sup = str(idx)
        if a.get('corresponding'):
            sup += ',' + corr_marker
        orcid = a.get('orcid', '')
        orcid_part = r'\,\orcidlink{' + orcid + '}' if orcid else ''
        # Title (Ünvan) — opsiyonel, varsa ad-soyad'ın önüne eklenir
        title = a.get('title', '').strip()
        name_with_title = (escape(title) + r'~' + escape(a['name'])) if title else escape(a['name'])
        parts.append(name_with_title + r'\textsuperscript{' + sup + r'}' + orcid_part)
    return ',\n  '.join(parts)


def _format_affiliations(authors: list, english_only: bool = False) -> str:
    lines = []
    email_label = 'E-mail: ' if english_only else 'E-posta: '
    for idx, a in enumerate(authors, 1):
        orcid = a.get('orcid', '')
        email = a.get('email', '')
        orcid_part = r' ORCID: \href{https://orcid.org/' + orcid + r'}{\mbox{' + orcid + r'}}.' if orcid else ''
        email_part = (r' ' + email_label + r'\href{mailto:' + email + r'}{\mbox{' + email + r'}}') if email else ''
        lines.append(
            r'\textsuperscript{' + str(idx) + r'}' +
            escape(a.get('affiliation', '')) + '.' +
            orcid_part + email_part
        )
    return r'\\' + '\n  '.join(lines)


def _format_corresponding(authors: list, english_only: bool = False) -> str:
    email_label = 'E-mail: ' if english_only else 'E-posta: '
    for a in authors:
        if a.get('corresponding'):
            name  = escape(a.get('name', ''))
            email = a.get('email', '')
            aff   = escape(a.get('affiliation', ''))
            ep    = r'. ' + email_label + r'\href{mailto:' + email + r'}{' + email + r'}' if email else ''
            return name + (', ' + aff if aff else '') + ep
    if authors:
        a = authors[0]
        name  = escape(a.get('name', ''))
        email = a.get('email', '')
        ep    = r'. ' + email_label + r'\href{mailto:' + email + r'}{' + email + r'}' if email else ''
        return name + ep
    return 'Author Name, Institution' if english_only else 'Yazar Adı, Kurum, E-posta'


# ── First-page LaTeX builder helpers ─────────────────────────────────────────
def _build_meta_strip(english_only: bool) -> str:
    """Sayfanın üst kısmındaki Yıl/Cilt/Sayı şeridi."""
    if english_only:
        return (r'{\fontsize{8}{10}\selectfont\quad Year: \SAARJyear\quad '
                r'Volume: \SAARJvolume\quad Issue: \SAARJissue\quad}')
    return (r'{\fontsize{8}{10}\selectfont\quad Year: \SAARJyear\quad '
            r'Volume: \SAARJvolume\quad Issue: \SAARJissue%'
            '\n       '
            r'\hfill Yıl: \SAARJyear\quad Cilt: \SAARJvolume\quad '
            r'Sayı: \SAARJissue\quad}')


def _build_titles(english_only: bool) -> str:
    """Türkçe ve İngilizce başlık satırları."""
    if english_only:
        return r'{\noindent{\fontsize{13}{15.5}\selectfont\bfseries \SAARJenglishtitle}\par}%'
    return (r'{\noindent{\fontsize{13}{15.5}\selectfont\bfseries \SAARJturkishtitle}\par}%'
            '\n\n  \\vspace{0.8mm}%\n\n  '
            r'{\noindent{\fontsize{11.5}{14}\selectfont\bfseries\itshape \SAARJenglishtitle}\par}%')


def _build_corresponding_label(english_only: bool, marker: str) -> str:
    """Dipnot satırındaki sorumlu yazar etiketi."""
    if english_only:
        return marker + r' Corresponding author'
    return marker + r' Sorumlu yazar / Corresponding author'


def _build_editor_row(editor: str, english_only: bool) -> str:
    """Sol sütunda anahtar kelimelerin altına eklenen opsiyonel 'İlgilenen Editör' satırı."""
    if not editor:
        return ''
    label = 'Editor:' if english_only else 'Editör / Editor:'
    return (
        '\n'
        r'      \vspace{2pt}\inforule%' + '\n'
        r'      \infoboldlabel{' + label + r'}\par\inforule%' + '\n'
        r'      \infovalue{' + escape(editor) + r'}\vspace{2pt}%'
    )


def _build_abstract_block(english_only: bool, has_tr: bool, has_en: bool, editor: str = '') -> str:
    """Özet/Abstract iki sütunlu blok. Boşsa ilgili sütunu gizler.

    English-only modunda yalnızca İngilizce blok gösterilir.
    """
    editor_row = _build_editor_row(editor, english_only)
    # English-only: tek sütun, sadece İngilizce
    if english_only:
        if not has_en:
            return ''
        return (
            r'\noindent\begin{minipage}[t]{\textwidth}%' + '\n'
            r'  \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%' + '\n'
            r'  \begin{tabular}{@{} p{0.183\textwidth} @{\hspace{2pt}} p{0.797\textwidth} @{}}%' + '\n'
            r'    \begin{minipage}[t]{0.183\textwidth}%' + '\n'
            r'      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%' + '\n'
            r'      \centering\infolabel{ARTICLE INFO}\par\inforule%' + '\n'
            r'      \infosubheading{Background:}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Received: \SAARJreceived}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Accepted: \SAARJaccepted}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Published: \SAARJpublished}\par\inforule%' + '\n'
            r'      \infoboldlabel{Keywords:}\par\inforule%' + '\n'
            r'      \infovalue{\SAARJenglishkeywords}\vspace{2pt}%' + editor_row + '\n'
            r'    \end{minipage}%' + '\n'
            r'    &%' + '\n'
            r'    \begin{minipage}[t]{0.797\textwidth}%' + '\n'
            r'      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%' + '\n'
            r'      {\fontsize{8.5}{10.5}\selectfont\bfseries\scshape Abstract}\par%' + '\n'
            r'      \noindent\rule{\linewidth}{0.4pt}\par%' + '\n'
            r'      {\fontsize{8.5}{10.5}\selectfont\SAARJenglishabstract}\vspace{2pt}%' + '\n'
            r'    \end{minipage}%' + '\n'
            r'  \end{tabular}%' + '\n'
            r'\end{minipage}%'
        )

    # İki dilli — TR ve/veya EN bloklarını koşullu çıkar
    blocks = []

    if has_tr:
        blocks.append(
            r'    \begin{minipage}[t]{0.183\textwidth}%' + '\n'
            r'      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%' + '\n'
            r'      \centering\infolabel{MAKALE BİLGİSİ}\par\inforule%' + '\n'
            r'      \infosubheading{Makale Geçmişi:}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Baş. tarihi: \SAARJreceived}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Kabul tarihi: \SAARJaccepted}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Yayın tarihi: \SAARJpublished}\par\inforule%' + '\n'
            r'      \infoboldlabel{Anahtar Kelimeler:}\par\inforule%' + '\n'
            r'      \infovalue{\SAARJturkishkeywords}\vspace{2pt}%' + editor_row + '\n'
            r'    \end{minipage}%' + '\n'
            r'    &%' + '\n'
            r'    \begin{minipage}[t]{0.797\textwidth}%' + '\n'
            r'      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%' + '\n'
            r'      {\fontsize{8.5}{10.5}\selectfont\bfseries\scshape Öz}\par%' + '\n'
            r'      \noindent\rule{\linewidth}{0.4pt}\par%' + '\n'
            r'      {\fontsize{8.5}{10.5}\selectfont\SAARJturkishabstract}\vspace{2pt}%' + '\n'
            r'    \end{minipage}%'
        )

    if has_en:
        blocks.append(
            r'    \begin{minipage}[t]{0.183\textwidth}%' + '\n'
            r'      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%' + '\n'
            r'      \centering\infolabel{ARTICLE INFO}\par\inforule%' + '\n'
            r'      \infosubheading{Background:}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Received: \SAARJreceived}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Accepted: \SAARJaccepted}\par\inforule%' + '\n'
            r'      {\fontsize{7.5}{9}\selectfont Published: \SAARJpublished}\par\inforule%' + '\n'
            r'      \infoboldlabel{Keywords:}\par\inforule%' + '\n'
            r'      \infovalue{\SAARJenglishkeywords}\vspace{2pt}%' + editor_row + '\n'
            r'    \end{minipage}%' + '\n'
            r'    &%' + '\n'
            r'    \begin{minipage}[t]{0.797\textwidth}%' + '\n'
            r'      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}\vspace{0pt}%' + '\n'
            r'      {\fontsize{8.5}{10.5}\selectfont\bfseries\scshape Abstract}\par%' + '\n'
            r'      \noindent\rule{\linewidth}{0.4pt}\par%' + '\n'
            r'      {\fontsize{8.5}{10.5}\selectfont\SAARJenglishabstract}\vspace{2pt}%' + '\n'
            r'    \end{minipage}%'
        )

    if not blocks:
        return ''

    sep = (r'    \\[3pt]%' + '\n'
           r'    \multicolumn{2}{@{}l@{}}{\rule{\dimexpr0.183\textwidth+0.797\textwidth+8pt\relax}{0.4pt}}\\[2pt]%' + '\n')

    body = blocks[0]
    for b in blocks[1:]:
        body += '\n' + r'    \\[3pt]%' + '\n' + sep + b
    body += r'\\[0pt]%'

    return (
        r'\noindent%' + '\n'
        r'\begin{tabular}{@{} p{0.183\textwidth} @{\hspace{2pt}} p{0.797\textwidth} @{}}%' + '\n'
        + body + '\n'
        r'\end{tabular}%'
    )


def _build_extended_abstract_section(ext: dict | None) -> str:
    if not ext:
        return ''

    fields = [
        ('English Title', ext.get('en_title', '').strip()),
        ('Background', ext.get('background', '').strip()),
        ('Research Purpose', ext.get('research_purpose', '').strip()),
        ('Methodology', ext.get('methodology', '').strip()),
        ('Findings', ext.get('findings', '').strip()),
        ('Conclusions', ext.get('conclusions', '').strip()),
    ]
    if not any(value for _, value in fields):
        return ''

    lines = [r'\section*{Extended Abstract}', '']
    for label, value in fields:
        if not value:
            continue
        if label == 'English Title':
            lines.append(r'{\noindent\bfseries ' + escape(value) + r'\par}')
        else:
            lines.append(r'\subsection*{' + label + '}')
            lines.append(render_rich_text(value))
        lines.append('')
    return '\n'.join(lines).strip()


def generate_latex(content: dict, authors: list, meta: dict) -> str:
    """Generate complete SAARJ-formatted LaTeX source."""

    tr_title   = content.get('tr_title') or r'Makalenin Türkçe Adı'
    en_title   = content.get('en_title') or r'Article Title in English'
    tr_abs     = content.get('tr_abstract') or r'Türkçe özet buraya yazılmalıdır.'
    en_abs     = content.get('en_abstract') or r'English abstract goes here.'
    tr_kw      = content.get('tr_keywords') or r'anahtar kelime 1; anahtar kelime 2'
    en_kw      = content.get('en_keywords') or r'keyword 1; keyword 2'

    year       = escape(meta.get('year', '2026'))
    volume     = escape(meta.get('volume', 'x'))
    issue      = escape(meta.get('issue', 'x'))
    start_page = escape(meta.get('start_page', 'xxx'))
    end_page   = escape(meta.get('end_page', 'xxx'))
    doi        = meta.get('doi', '')

    author_short = meta.get('author_short', '')
    if not author_short and authors:
        names = [a['name'].split()[-1] for a in authors]
        if len(names) == 1:
            author_short = escape(names[0])
        elif len(names) == 2:
            author_short = escape(names[0]) + r' \& ' + escape(names[1])
        else:
            author_short = escape(names[0]) + r' et al.'

    head_title = meta.get('head_title', '') or tr_title[:60]

    author_block  = _format_author_block(authors) if authors else r'Author One\textsuperscript{1,*}'
    affiliations  = _format_affiliations(authors) if authors else r'\textsuperscript{1}Kurum Adı, Şehir.'
    corresponding = _format_corresponding(authors) if authors else r'Yazar Adı, Kurum.'

    # ── body sections ──
    body_lines = []
    for sec in content.get('sections', []):
        title = sec['title']
        level = sec.get('level', 1)
        cmd   = SECTION_CMD.get(level, r'\section')
        body_lines.append(cmd + '{' + title + '}')
        if sec['latex']:
            body_lines.append('')
            body_lines.append(sec['latex'])
        body_lines.append('')

    # ── references ──
    refs_tex = ''
    raw_refs = content.get('references', [])
    if raw_refs:
        ref_items = []
        for idx, r in enumerate(raw_refs, 1):
            ref_items.append(r'  \bibitem{ref' + str(idx) + r'}' + '\n  ' + r)
        refs_tex = (
            r'\section*{Kaynakça / References}' + '\n'
            r'\renewcommand{\refname}{}' + '\n'
            r'\vspace{-2\baselineskip}' + '\n'
            r'\begin{thebibliography}{99}' + '\n'
            r'\setlength{\leftmargin}{1.5em}%' + '\n'
            r'\setlength{\itemindent}{-1.5em}%' + '\n\n' +
            '\n\n'.join(ref_items) + '\n\n'
            r'\end{thebibliography}'
        )
    else:
        refs_tex = (
            r'\section*{Kaynakça / References}' + '\n'
            r'\renewcommand{\refname}{}' + '\n'
            r'\vspace{-2\baselineskip}' + '\n'
            r'\begin{thebibliography}{99}' + '\n'
            r'\setlength{\leftmargin}{1.5em}%' + '\n'
            r'\setlength{\itemindent}{-1.5em}%' + '\n\n'
            r'% Kaynakları buraya ekleyin / Add your references here' + '\n\n'
            r'\end{thebibliography}'
        )

    apa_citation = (
        author_short + r'\ (' + year + r'). ' +
        tr_title + r'. \textit{Samsun Aviation and Aeronautical Research Journal}, ' +
        r'\textit{' + volume + r'}(' + issue + r'), ' +
        start_page + r'--' + end_page + r'.'
    )

    body_text = '\n'.join(body_lines) if body_lines else (
        r'\section{Giriş / Introduction}' + '\n\n'
        r'% Makale metni buraya gelecek / Article body goes here' + '\n'
    )

    doi_line = doi if doi else ''

    tex = r"""% ============================================================
%  SAARJ — Samsun Aviation and Aeronautical Research Journal
%  Bu dosya SAARJ Template uygulaması tarafından otomatik oluşturulmuştur.
%  Overleaf'e yükleyin: SAARJ.png ile birlikte ZIP'i açın.
%  Compile with: XeLaTeX
% ============================================================

\documentclass[10pt,a4paper]{article}

% --- Font & Unicode ---
\usepackage{fontspec}
\usepackage{unicode-math}

% --- Layout ---
\usepackage{geometry}
\usepackage{fancyhdr}
\usepackage{multicol}

% --- Typography ---
\usepackage{microtype}
\usepackage{xcolor}
\usepackage{relsize}

% --- Section headings ---
\usepackage{titlesec}

% --- Tables ---
\usepackage{array}
\usepackage{tabularx}
\usepackage{booktabs}
\usepackage{multirow}
\usepackage{makecell}
\usepackage{longtable}

% --- Figures ---
\usepackage{graphicx}
\usepackage{float}
\usepackage{caption}

% --- Mathematics ---
\usepackage{amsmath}

% --- References & Links ---
\usepackage{hyperref}
\usepackage{url}
\usepackage{doi}
\usepackage[numbers,sort&compress]{natbib}
\usepackage{orcidlink}

% --- Footnotes / Lists ---
\usepackage{footmisc}
\usepackage{enumitem}
\usepackage{hanging}

% --- Misc ---
\usepackage{etoolbox}
\usepackage{calc}
\usepackage{lastpage}
\usepackage{ifthen}


% ============================================================
%  FONT
% ============================================================
\setmainfont{TeX Gyre Pagella}
\setmathfont{TeX Gyre Pagella Math}


% ============================================================
%  COLOURS
% ============================================================
\definecolor{SAARJbrown}{HTML}{833C0B}
\definecolor{SAARJblue}{HTML}{0070C0}
\definecolor{SAARJgray}{HTML}{CFCDCD}
\definecolor{SAARJdarkgray}{HTML}{3B3838}


% ============================================================
%  HYPERLINKS
% ============================================================
\hypersetup{
  colorlinks=true, urlcolor=SAARJblue,
  linkcolor=black, citecolor=black,
  pdfencoding=auto, unicode=true,
}


% ============================================================
%  GEOMETRY (body pages)
% ============================================================
\geometry{
  a4paper,
  top=1.5cm, bottom=1.5cm, left=1.5cm, right=1.5cm,
  headheight=1.2cm, headsep=0.4cm, footskip=0.8cm,
}


% ============================================================
%  PARAGRAPH FORMAT
% ============================================================
\setlength{\parindent}{0pt}
\setlength{\parskip}{4pt}
\renewcommand{\baselinestretch}{1.0}


% ============================================================
%  SECTION HEADINGS
% ============================================================
\titleformat{\section}[block]{\fontsize{11}{13}\selectfont\bfseries\centering}{}{0em}{}
\titlespacing*{\section}{0pt}{9pt}{5pt}
\titleformat*{\section}{\fontsize{11}{13}\selectfont\bfseries\centering}

\titleformat{\subsection}[block]{\fontsize{11}{13}\selectfont\bfseries}{}{0em}{}
\titlespacing*{\subsection}{0pt}{7pt}{3pt}

\titleformat{\subsubsection}[block]{\fontsize{11}{13}\selectfont\bfseries\itshape}{}{0em}{}
\titlespacing*{\subsubsection}{0pt}{6pt}{3pt}


% ============================================================
%  HEADERS & FOOTERS
% ============================================================
\pagestyle{fancy}
\fancyhf{}
\fancyhead[C]{%
  \fontsize{8.5}{10.5}\selectfont\itshape
  """ + author_short + r"""\ (""" + year + r""").
  """ + head_title + r""".
  \textup{Samsun Aviation and Aeronautical Research Journal},
  \textit{""" + volume + r"""}(""" + issue + r"""),
  """ + start_page + r"""--""" + end_page + r"""%
}
\fancyfoot[C]{\fontsize{9}{11}\selectfont\thepage}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

\fancypagestyle{firstpage}{%
  \fancyhf{}%
  \renewcommand{\headrulewidth}{0pt}%
}


% ============================================================
%  CAPTION FORMAT
% ============================================================
\captionsetup{font={small,bf}, labelsep=period, justification=centering, skip=4pt}
\captionsetup[table]{position=top}


% ============================================================
%  FOOTNOTE FORMAT
% ============================================================
\renewcommand{\footnoterule}{\kern-3pt\hrule width 2cm height 0.4pt\kern 2.6pt}
\setlength{\footnotesep}{4pt}


% ============================================================
%  LIST FORMAT
% ============================================================
\setlist{leftmargin=1.2cm, itemsep=0pt, parsep=0pt, topsep=2pt, partopsep=0pt}


% ============================================================
%  ARTICLE METADATA
% ============================================================
\newcommand{\SAARJyear}{""" + year + r"""}
\newcommand{\SAARJvolume}{""" + volume + r"""}
\newcommand{\SAARJissue}{""" + issue + r"""}
\newcommand{\SAARJstartpage}{""" + start_page + r"""}
\newcommand{\SAARJendpage}{""" + end_page + r"""}
\newcommand{\SAARJDOI}{""" + doi_line + r"""}

\newcommand{\SAARJarticletype}{Araştırma Makalesi -- Research Article}
\newcommand{\SAARJreceived}{xx.xx.xxxx}
\newcommand{\SAARJaccepted}{xx.xx.xxxx}
\newcommand{\SAARJpublished}{xx.xx.xxxx}

\newcommand{\SAARJturkishtitle}{""" + tr_title + r"""}
\newcommand{\SAARJturkishabstract}{""" + escape(tr_abs) + r"""}
\newcommand{\SAARJturkishkeywords}{""" + tr_kw + r"""}

\newcommand{\SAARJenglishtitle}{""" + en_title + r"""}
\newcommand{\SAARJenglishabstract}{""" + escape(en_abs) + r"""}
\newcommand{\SAARJenglishkeywords}{""" + en_kw + r"""}

\newcommand{\SAARJauthorshort}{""" + author_short + r"""}
\newcommand{\SAARJheadtitle}{""" + head_title + r"""}

\newcommand{\SAARJaffiliations}{""" + affiliations + r"""}
\newcommand{\SAARJcorrespondinginfo}{""" + corresponding + r"""}

\newcommand{\SAARJapacitation}{""" + apa_citation + r"""}
\newcommand{\SAARJethicsstatement}{%
  Bu araştırma, ilgili etik kurul kararı doğrultusunda yürütülmüş olup
  tüm etik ilkelere uyulmuştur. / This study was conducted in accordance
  with the relevant ethics committee decision and all ethical principles
  were followed.%
}


% ============================================================
%  INTERNAL HELPERS
% ============================================================
\newcommand{\infolabel}[1]{{\fontsize{9}{11}\selectfont\scshape #1}}
\newcommand{\infosubheading}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\infovalue}[1]{{\fontsize{7.5}{9}\selectfont #1}}
\newcommand{\infoboldlabel}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\inforule}{%
  \par\vspace{0pt}%
  \noindent\rule{\linewidth}{0.4pt}%
  \par\vspace{0pt}%
}

% ============================================================
%  BIBLIOGRAPHY
% ============================================================
\bibliographystyle{unsrtnat}
\setlength{\bibhang}{1.5em}
\setlength{\bibsep}{3pt}


% ============================================================
%  FIRST PAGE COMMAND
% ============================================================
\newcommand{\SAARJfirstpage}[1]{%

  \newgeometry{
    a4paper,
    top=1.2cm, bottom=1.5cm, left=1.5cm, right=1.5cm,
    headheight=0pt, headsep=0pt, footskip=0.8cm,
  }%
  \thispagestyle{firstpage}%

  % --- Logo | ISSN / URL / DOI ---
  \noindent%
  {\setlength{\tabcolsep}{0pt}%
  \setlength{\extrarowheight}{0pt}%
  \begin{tabular}{@{} m{11cm} @{} >{\raggedleft\arraybackslash}m{7cm} @{}}%
    \includegraphics[height=2.3cm]{SAARJ}%
    &%
    \raggedleft%
    {\fontsize{8.5}{11}\selectfont\color{SAARJdarkgray}%
      \textbf{ISSN:} 2717-6924\\[2pt]%
      \href{https://dergipark.org.tr/en/pub/saarj}{%
        \color{SAARJblue}\itshape https://dergipark.org.tr/en/pub/saarj%
      }\\[2pt]%
      \ifthenelse{\equal{\SAARJDOI}{}}{%
        \textbf{DOI:} \textit{(atanacak\,/\,to be assigned)}%
      }{%
        \textbf{DOI:} \doi{\SAARJDOI}%
      }\par%
    }%
  \end{tabular}}%

  \vspace{-1.5mm}%
  \noindent\rule{\textwidth}{0.4pt}%

  {\noindent%
  \setlength{\fboxsep}{1.5pt}\setlength{\fboxrule}{0pt}%
  \colorbox{SAARJgray}{%
    \begin{minipage}{\dimexpr\textwidth-3pt\relax}%
      {\fontsize{8}{10}\selectfont%
        \quad Year: \SAARJyear\quad Volume: \SAARJvolume\quad Issue: \SAARJissue%
        \hfill%
        Yıl: \SAARJyear\quad Cilt: \SAARJvolume\quad Sayı: \SAARJissue\quad%
      }%
    \end{minipage}%
  }}%

  \vspace{0.8mm}%

  {\noindent%
  \fontsize{8}{10}\selectfont\color{SAARJdarkgray}%
  \ifthenelse{\equal{\SAARJDOI}{}}{%
    \textbf{DOI:} \textit{(atanacak\,/\,to be assigned)}%
  }{%
    \textbf{DOI:} \doi{\SAARJDOI}%
  }\par}%

  \vspace{1mm}%

  {\noindent\centering%
  {\fontsize{10.5}{12}\selectfont\bfseries\itshape \SAARJarticletype}\par}%

  \vspace{1.5mm}%

  {\noindent{\fontsize{13}{15.5}\selectfont\bfseries \SAARJturkishtitle}\par}%

  \vspace{0.8mm}%

  {\noindent{\fontsize{11.5}{14}\selectfont\bfseries\itshape \SAARJenglishtitle}\par}%

  \vspace{1.5mm}%

  {\noindent{\fontsize{10}{12}\selectfont #1}\par}%

  \vspace{0.5mm}%

  \setlength{\tabcolsep}{4pt}%
  \noindent\rule{\textwidth}{0.4pt}%

  \noindent%
  \begin{tabular}{@{} p{0.183\textwidth} @{\hspace{2pt}} p{0.797\textwidth} @{}}%

    \begin{minipage}[t]{0.183\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      \centering\infolabel{MAKALE BİLGİSİ}\par%
      \inforule%
      \infosubheading{Makale Geçmişi:}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Baş. tarihi: \SAARJreceived}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Kabul tarihi: \SAARJaccepted}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Yayın tarihi: \SAARJpublished}\par%
      \inforule%
      \infoboldlabel{Anahtar Kelimeler:}\par%
      \inforule%
      \infovalue{\SAARJturkishkeywords}%
      \vspace{2pt}%
    \end{minipage}%
    &%
    \begin{minipage}[t]{0.797\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      {\fontsize{9}{11}\selectfont\bfseries\scshape Özet}\par%
      \noindent\rule{\linewidth}{0.4pt}\par%
      {\fontsize{9}{11}\selectfont\SAARJturkishabstract}%
      \vspace{2pt}%
    \end{minipage}%
    \\[3pt]%

    \multicolumn{2}{@{}l@{}}{\rule{\dimexpr0.183\textwidth+0.797\textwidth+8pt\relax}{0.4pt}}\\[2pt]%

    \begin{minipage}[t]{0.183\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      \centering\infolabel{ARTICLE INFO}\par%
      \inforule%
      \infosubheading{Background:}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Received: \SAARJreceived}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Accepted: \SAARJaccepted}\par%
      \inforule%
      {\fontsize{7.5}{9}\selectfont Published: \SAARJpublished}\par%
      \inforule%
      \infoboldlabel{Keywords:}\par%
      \inforule%
      \infovalue{\SAARJenglishkeywords}%
      \vspace{2pt}%
    \end{minipage}%
    &%
    \begin{minipage}[t]{0.797\textwidth}%
      \setlength{\parindent}{0pt}\setlength{\parskip}{0pt}%
      \vspace{0pt}%
      {\fontsize{9}{11}\selectfont\bfseries\scshape Abstract}\par%
      \noindent\rule{\linewidth}{0.4pt}\par%
      {\fontsize{9}{11}\selectfont\SAARJenglishabstract}%
      \vspace{2pt}%
    \end{minipage}%
    \\[0pt]%

  \end{tabular}%

  \noindent\rule{\textwidth}{0.4pt}%

  \vspace*{\fill}%
  \noindent\begin{minipage}{\textwidth}%
  \begingroup
    \sloppy\emergencystretch=3em%
    \setlength{\parskip}{1pt}\setlength{\parindent}{0pt}%
    \fontsize{7}{8.5}\selectfont%
    \noindent\rule{\textwidth}{0.4pt}\par%
    \noindent\SAARJaffiliations\par%
    \noindent\rule{\textwidth}{0.2pt}\par%
    \noindent\textit{*Sorumlu yazar / Corresponding author}\par%
    \noindent\textbf{Önerilen Atıf / Suggested Citation:} \SAARJapacitation\par%
    \noindent{\setlength{\parskip}{0pt}\textbf{Etik Beyan / Ethics Statement:} \SAARJethicsstatement\par}%
  \endgroup%
  \end{minipage}

  \restoregeometry%
}% end \SAARJfirstpage


% ============================================================
%  BEGIN DOCUMENT
% ============================================================
\begin{document}

\SAARJfirstpage{%
  """ + author_block + r"""%
}
% Kapak sayfası sayfa 1'dir ancak numara gösterilmez.
% İkinci sayfa sayfa 2 olarak başlar.

""" + body_text + r"""

\section*{Araştırmacıların Katkı Oranı / Author Contributions}
Kavramsal çerçeve / Conceptualization: ; Yöntem / Methodology: ;
Veri toplama / Data collection: ; Analiz / Analysis: ;
Yazım / Writing: ; Gözden geçirme / Review \& Editing:

\section*{Çıkar Çatışması / Conflict of Interest}
Yazarlar herhangi bir çıkar çatışması olmadığını beyan eder. /
The authors declare no conflict of interest.

""" + refs_tex + r"""

\end{document}
"""
    return tex


# ── ZIP builder ────────────────────────────────────────────────────────────────
def build_zip(tex_content: str, logo_src: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('main.tex', tex_content.encode('utf-8'))

        if os.path.exists(logo_src):
            zf.write(logo_src, 'SAARJ.png')

        readme = """SAARJ LaTeX Formatter — Overleaf Yükleme Rehberi
=================================================

1. Bu ZIP dosyasını açın.
2. Overleaf.com adresine gidin → New Project → Upload Project
3. ZIP dosyasının tamamını yükleyin (main.tex + SAARJ.png).
4. Derleyici olarak XeLaTeX seçin:
   Menu → Compiler → XeLaTeX
5. "Recompile" butonuna tıklayın.

Düzenleme önerileri:
- Metadata (yıl, cilt, sayı, sayfalar): dosyanın üst kısmındaki \\newcommand satırları
- Makale tarihleri: \\SAARJreceived, \\SAARJaccepted, \\SAARJpublished
- Etik beyan: \\SAARJethicsstatement
- Katkı oranları: Araştırmacıların Katkı Oranı bölümü

Sorular için: SAARJ Template uygulamasına bakın.
"""
        zf.writestr('README_Overleaf.txt', readme.encode('utf-8'))

        author_template = """# SAARJ Yazar Bilgileri Şablonu
# Her satır bir yazar — sütunlar: | ile ayrılır
# Sütunlar: Ad Soyad | Kurum | ORCID | E-posta | Sorumlu (evet/hayır)
# Örnek:

Ahmet Yılmaz | Turizm Bölümü, Ankara Üniversitesi, Ankara | 0000-0000-0000-0001 | ahmet@uni.edu.tr | evet
Ayşe Kaya | İşletme Bölümü, İstanbul Üniversitesi, İstanbul | 0000-0000-0000-0002 | ayse@uni.edu.tr | hayır
"""
        zf.writestr('yazar_bilgileri_sablonu.txt', author_template.encode('utf-8'))
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
#  FORM-BASED LaTeX GENERATION
# ══════════════════════════════════════════════════════════════════════════════

SECTION_CMD_MAP = {'1': r'\section', '2': r'\subsection', '3': r'\subsubsection'}

# Special starred sections (no numbering)
STARRED_NAMES = {
    'araştırmacı', 'katkı', 'contributions', 'çıkar', 'conflict',
    'teşekkür', 'acknowledgement', 'kaynakça', 'references', 'bibliography',
}


def _table_to_latex(text: str) -> str:
    """Convert pipe-separated plain text to a LaTeX tabular body."""
    lines = [l.strip() for l in text.strip().splitlines() if l.strip()]
    if not lines:
        return ''
    rows = [[escape(c.strip()) for c in l.split('|')] for l in lines]
    ncols = max(len(r) for r in rows)
    col_spec = 'l' * ncols
    out = [r'\begin{tabular}{@{}' + col_spec + r'@{}}', r'\toprule']
    for i, row in enumerate(rows):
        # pad short rows
        while len(row) < ncols:
            row.append('')
        out.append(' & '.join(row) + r' \\')
        if i == 0:
            out.append(r'\midrule')
    out.append(r'\bottomrule')
    out.append(r'\end{tabular}')
    return '\n'.join(out)


def _build_figtable_latex(ft: dict, file_ext_map: dict, article_language: str = 'turkish') -> str:
    """Return LaTeX for one figure or table item."""
    num      = ft.get('number', '1')
    tr_cap   = escape(ft.get('tr_cap', ''))
    en_cap   = escape(ft.get('en_cap', ''))
    if article_language == 'english':
        caption = en_cap or tr_cap
    else:
        caption = tr_cap or en_cap
    label_prefix = 'fig' if ft['type'] == 'figure' else 'tab'
    label    = label_prefix + ':' + re.sub(r'\W+', '_', num)

    if ft['type'] == 'figure':
        fkey = ft.get('file_key', '')
        fname = file_ext_map.get(fkey, 'fig_' + fkey)
        return (
            r'\begin{figure}[htbp]' + '\n'
            r'  \centering' + '\n'
            r'  \includegraphics[width=\linewidth]{' + fname + '}\n'
            r'  \caption{' + caption + '}\n'
            r'  \label{' + label + '}\n'
            r'\end{figure}'
        )
    else:
        tbl_body = _table_to_latex(ft.get('tbl_data', ''))
        if not tbl_body:
            tbl_body = '% Tablo içeriği buraya gelecek / Table content here'
        return (
            r'\begin{table}[htbp]' + '\n'
            r'  \centering' + '\n'
            r'  \caption{' + caption + '}\n'
            r'  \label{' + label + '}\n'
            '  ' + tbl_body.replace('\n', '\n  ') + '\n'
            r'\end{table}'
        )


def generate_latex_from_form(data: dict, figure_file_bytes: dict,
                             journal_settings: dict = None) -> str:
    """
    Generate complete LaTeX from structured form data.
    data keys: cover, authors, abstract, sections, figtables, extra, references
    figure_file_bytes: {file_key: (filename_in_zip, bytes)} for figure files
    journal_settings: optional journal branding/typography overrides
    """
    # ── Journal settings (with SAARJ defaults) ──
    js          = journal_settings or {}
    jname_en    = js.get('journal_name_en',  'Samsun Aviation and Aeronautical Research Journal')
    jname_tr    = js.get('journal_name_tr',  'Samsun Havacılık Araştırmaları Dergisi')
    issn_print  = js.get('issn_print',  '2717-6924')
    issn_online = js.get('issn_online', '')
    j_url       = js.get('journal_url', 'https://dergipark.org.tr/en/pub/saarj')
    font_name   = js.get('font',        'texgyrepagella')
    body_size   = js.get('body_size',   '10')
    accent_hex  = js.get('accent_color', '#833C0B').lstrip('#')
    logo_stem   = js.get('logo_stem',   'journal_logo')
    # Yeni opsiyonlar
    corr_marker       = js.get('corresponding_marker', '*') or '*'
    cc_logo_stem      = js.get('cc_logo_stem', 'ccby')
    logo_height_cm    = float(js.get('logo_height_cm', 2.3) or 2.3)
    # DOI konumu: 'top' (URL altı) veya 'bottom' (sağ alt). Varsayılan 'bottom'.
    doi_position      = (js.get('doi_position', 'bottom') or 'bottom').strip().lower()
    if doi_position not in ('top', 'bottom'):
        doi_position = 'bottom'
    # Top margin'i logo yüksekliğine göre ayarla (logo büyüdükçe üst boşluk azalır)
    if logo_height_cm <= 1.5:
        top_margin_cm = 1.2
    elif logo_height_cm <= 2.0:
        top_margin_cm = 1.0
    elif logo_height_cm <= 2.6:
        top_margin_cm = 0.9
    elif logo_height_cm <= 3.2:
        top_margin_cm = 0.7
    else:
        top_margin_cm = 0.5

    # Font setup in LaTeX
    # Kullanıcı dostu ad → (Overleaf/TeX Live display adı, math font adı veya None)
    _FONT_MAP = {
        'palatino linotype':  ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'palatino':           ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'texgyrepagella':     ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'tex gyre pagella':   ('TeX Gyre Pagella',   'TeX Gyre Pagella Math'),
        'times new roman':    ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'times':              ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'texgyretermes':      ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'tex gyre termes':    ('TeX Gyre Termes',    'TeX Gyre Termes Math'),
        'century':            ('TeX Gyre Bonum',     None),
        'century schoolbook': ('TeX Gyre Bonum',     None),
        'texgyrebonum':       ('TeX Gyre Bonum',     None),
        'calibri':            ('Carlito',            None),
        'sans serif':         ('TeX Gyre Heros',     None),
        'sans-serif':         ('TeX Gyre Heros',     None),
        'arial':              ('TeX Gyre Heros',     None),
        'texgyreheros':       ('TeX Gyre Heros',     None),
        'latinmodern':        ('Latin Modern Roman', 'Latin Modern Math'),
        'latin modern':       ('Latin Modern Roman', 'Latin Modern Math'),
    }
    _key = font_name.lower().strip()
    if _key in _FONT_MAP:
        _display, _mathfont = _FONT_MAP[_key]
        font_setup = r'\setmainfont{' + _display + '}'
        if _mathfont:
            font_setup += '\n' + r'\setmathfont{' + _mathfont + '}'
    else:
        font_setup = r'\setmainfont{' + font_name + '}'

    # ISSN display — yalnızca girilmiş alanları göster
    _issn_parts = []
    if issn_print:
        _issn_parts.append(r'\textbf{ISSN:} ' + issn_print)
    if issn_online:
        _issn_parts.append(r'\textbf{e-ISSN:} ' + issn_online)
    issn_display = (r'\ \ '.join(_issn_parts) + r'\\[2pt]%') if _issn_parts else ''

    cov      = data.get('cover', {})
    authors  = data.get('authors', [])
    abstr    = data.get('abstract', {})
    sections = data.get('sections', [])
    fts      = data.get('figtables', [])
    extra    = data.get('extra', {})
    refs_raw = data.get('references', '')
    ext_abs  = data.get('extended_abstract') or {}
    article_language = (cov.get('article_language', 'turkish') or 'turkish').strip().lower()
    english_only = article_language == 'english'

    # ── Cover fields ──
    tr_title   = escape(cov.get('tr_title', '') or 'Makalenin Türkçe Adı')
    en_title   = escape(cov.get('en_title', '') or 'Article Title in English')
    year       = escape(cov.get('year',       '2026'))
    volume     = escape(cov.get('volume',     'x'))
    issue      = escape(cov.get('issue',      'x'))
    start_page = escape(cov.get('start_page', 'xxx'))
    end_page   = escape(cov.get('end_page',   'xxx'))
    doi        = cov.get('doi', '')
    art_type   = escape(cov.get('article_type', 'Araştırma Makalesi -- Research Article'))
    received   = escape(cov.get('received',  'xx.xx.xxxx'))
    accepted   = escape(cov.get('accepted',  'xx.xx.xxxx'))
    published  = escape(cov.get('published', 'xx.xx.xxxx'))
    editor_raw = (cov.get('editor', '') or '').strip()
    ethics_raw = cov.get('ethics', '').strip()
    ethics     = escape(ethics_raw) if ethics_raw else (
        r'Bu araştırma, ilgili etik kurul kararı doğrultusunda yürütülmüş olup '
        r'tüm etik ilkelere uyulmuştur. / This study was conducted in accordance '
        r'with the relevant ethics committee decision and all ethical principles were followed.'
    )

    # ── Author short / head title ──
    author_short = cov.get('author_short', '').strip()
    if not author_short and authors:
        names = [a['name'].split()[-1] for a in authors if a.get('name')]
        if len(names) == 1:
            author_short = escape(names[0])
        elif len(names) == 2:
            author_short = escape(names[0]) + r' \& ' + escape(names[1])
        elif names:
            author_short = escape(names[0]) + r' et al.'
    if not author_short:
        author_short = 'Yazar'
    else:
        author_short = escape(author_short)

    display_title = en_title if english_only or not tr_title else tr_title
    head_title = display_title

    # ── Abstract / keywords (boş olabilirler) ──
    tr_abs_raw = (abstr.get('tr_abs', '') or '').strip()
    en_abs_raw = (abstr.get('en_abs', '') or '').strip()
    tr_kw_raw  = (abstr.get('tr_kw',  '') or '').strip()
    en_kw_raw  = (abstr.get('en_kw',  '') or '').strip()
    tr_abs = escape(tr_abs_raw)
    en_abs = escape(en_abs_raw)
    tr_kw  = escape(tr_kw_raw)
    en_kw  = escape(en_kw_raw)
    has_tr_abs = bool(tr_abs_raw) and not english_only
    has_en_abs = bool(en_abs_raw)

    # ── Author blocks ──
    author_block  = _format_author_block(authors, corr_marker)        if authors else r'Author One\textsuperscript{1,' + corr_marker + r'}'
    affiliations  = _format_affiliations(authors, english_only)       if authors else (r'\textsuperscript{1}Institution, City.' if english_only else r'\textsuperscript{1}Kurum Adı, Şehir.')
    corresponding = _format_corresponding(authors, english_only)      if authors else (r'Author Name, Institution' if english_only else r'Yazar Adı, Kurum.')

    # ── APA citation ──
    apa_citation = (
        author_short + r'\ (' + year + r'). ' + display_title +
        r'. \textit{' + escape(jname_en) + r'}, ' +
        r'\textit{' + volume + r'}(' + issue + r'), ' +
        start_page + r'--' + end_page + r'.'
    )

    # ── Figure file → zip name mapping ──
    # file_ext_map: {file_key: 'fig_N'} (no extension; XeLaTeX resolves)
    file_ext_map = {}
    for fkey, (zipname, _) in figure_file_bytes.items():
        # zipname like 'fig_3.png' → stem 'fig_3'
        stem = zipname.rsplit('.', 1)[0] if '.' in zipname else zipname
        file_ext_map[fkey] = stem

    # ── Body sections + paragraph-level figure/table placement ──
    body_lines = []
    placed_ft_ids = set()   # track which fts have already been placed

    for sec in sections:
        name    = sec.get('name', 'Bölüm')
        level   = sec.get('level', '1')
        content = sec.get('content', '').strip()
        cmd     = SECTION_CMD_MAP.get(level, r'\section')
        starred = any(k in name.lower() for k in STARRED_NAMES)
        star    = '*' if starred else ''
        body_lines.append(cmd + star + '{' + escape(name) + '}')
        body_lines.append('')

        # FTs assigned to this section
        sec_fts = [
            (i, ft) for i, ft in enumerate(fts)
            if ft.get('section', '').strip() == name.strip() and i not in placed_ft_ids
        ]

        if content:
            # Split content into paragraphs (blank-line separated or single newlines)
            raw_paras = [p.strip() for p in re.split(r'\n\s*\n', content)]
            raw_paras = [p for p in raw_paras if p]
            if not raw_paras:
                raw_paras = [content]

            for para in raw_paras:
                body_lines.append(render_rich_text(para))
                body_lines.append('')

                # Check if any ft's anchor text is found in this paragraph
                for i, ft in sec_fts:
                    if i in placed_ft_ids:
                        continue
                    anchor = ft.get('after_para', '').strip()
                    if anchor and anchor.lower() in para.lower():
                        body_lines.append(_build_figtable_latex(ft, file_ext_map, article_language))
                        body_lines.append('')
                        placed_ft_ids.add(i)

            # FTs for this section with anchor NOT found → append at section end
            for i, ft in sec_fts:
                if i not in placed_ft_ids:
                    body_lines.append(_build_figtable_latex(ft, file_ext_map, article_language))
                    body_lines.append('')
                    placed_ft_ids.add(i)
        else:
            # No content — place all section FTs here
            for i, ft in sec_fts:
                if i not in placed_ft_ids:
                    body_lines.append(_build_figtable_latex(ft, file_ext_map, article_language))
                    body_lines.append('')
                    placed_ft_ids.add(i)

    if not body_lines:
        body_lines = [
            r'\section{Giriş / Introduction}',
            '',
            r'% Makale metni buraya gelecek / Article body goes here',
            '',
        ]

    # ── Figures/tables with no section assigned (end of body) ──
    orphan_fts = [
        (i, ft) for i, ft in enumerate(fts)
        if not ft.get('section', '').strip() and i not in placed_ft_ids
    ]
    if orphan_fts:
        body_lines.append(r'% ── Şekil ve Tablolar / Figures and Tables ──')
        body_lines.append(r'\clearpage')
        for i, ft in orphan_fts:
            body_lines.append(_build_figtable_latex(ft, file_ext_map, article_language))
            body_lines.append('')
            placed_ft_ids.add(i)

    # ── Extra sections (başlıklar dil moduna göre) ──
    ack_heading      = 'Acknowledgements' if english_only else r'Teşekkür / Acknowledgements'
    contrib_heading  = 'Author Contributions' if english_only else r'Araştırmacıların Katkı Oranı / Author Contributions'
    conflict_heading = 'Conflict of Interest' if english_only else r'Çıkar Çatışması / Conflict of Interest'

    ack = extra.get('ack', '').strip()
    if ack:
        body_lines.append(r'\section*{' + ack_heading + r'}')
        body_lines.append(render_rich_text(ack))
        body_lines.append('')

    contrib = extra.get('contrib', '').strip()
    body_lines.append(r'\section*{' + contrib_heading + r'}')
    if contrib:
        body_lines.append(render_rich_text(contrib))
    elif english_only:
        body_lines.append(
            r'Conceptualization: ; Methodology: ; Data collection: ; '
            r'Analysis: ; Writing: ; Review \& Editing:'
        )
    else:
        body_lines.append(
            r'Kavramsal çerçeve / Conceptualization: ; Yöntem / Methodology: ; '
            r'Veri toplama / Data collection: ; Analiz / Analysis: ; '
            r'Yazım / Writing: ; Gözden geçirme / Review \& Editing:'
        )
    body_lines.append('')

    conflict = extra.get('conflict', '').strip()
    body_lines.append(r'\section*{' + conflict_heading + r'}')
    if conflict:
        body_lines.append(render_rich_text(conflict))
    elif english_only:
        body_lines.append(r'The authors declare no conflict of interest.')
    else:
        body_lines.append(
            r'Yazarlar herhangi bir çıkar çatışması olmadığını beyan eder. / '
            r'The authors declare no conflict of interest.'
        )
    body_lines.append('')

    # ── References — alphabetically sorted, hanging indent, no numbers ──
    refs_lines = sorted(
        [l.strip() for l in refs_raw.splitlines() if l.strip()],
        key=lambda x: x.lower()
    )
    refs_heading = 'References' if english_only else r'Kaynakça / References'
    _ref_env_open = (
        r'\section*{' + refs_heading + r'}' + '\n'
        r'\begin{list}{}{%' + '\n'
        r'  \setlength{\leftmargin}{1.5em}%' + '\n'
        r'  \setlength{\itemindent}{-1.5em}%' + '\n'
        r'  \setlength{\topsep}{2pt}%' + '\n'
        r'  \setlength{\itemsep}{3pt}%' + '\n'
        r'  \setlength{\parsep}{0pt}%' + '\n'
        r'}' + '\n'
    )
    if refs_lines:
        items = '\n'.join(r'\item ' + render_rich_text(r) for r in refs_lines)
        refs_tex = _ref_env_open + items + '\n' + r'\end{list}'
    else:
        refs_tex = (
            _ref_env_open +
            r'% Kaynakları buraya ekleyin / Add your references here' + '\n' +
            r'\end{list}'
        )

    extended_abstract_tex = ''
    if not english_only:
        extended_abstract_tex = _build_extended_abstract_section(ext_abs)

    body_text = '\n'.join(body_lines)
    doi_line  = doi if doi else ''

    # ── Full LaTeX document ──
    tex = r"""% ============================================================
%  """ + escape(jname_en) + r"""
%  Bu dosya Journal LaTeX Formatter tarafından oluşturulmuştur.
%  Overleaf: New Project → Upload Project → bu ZIP'i seçin.
%  Derleyici: XeLaTeX
% ============================================================

\documentclass[""" + body_size + r"""pt,a4paper]{article}

\usepackage{fontspec}
\usepackage{unicode-math}
\usepackage{geometry}
\usepackage{fancyhdr}
\usepackage{microtype}
\usepackage{xcolor}
\usepackage{titlesec}
\usepackage{array}
\usepackage{tabularx}
\usepackage{booktabs}
\usepackage{multirow}
\usepackage{makecell}
\usepackage{longtable}
\usepackage{graphicx}
\usepackage{float}
\usepackage{caption}
\usepackage{amsmath}
\usepackage{hyperref}
\usepackage{url}
\usepackage{doi}
\usepackage[numbers,sort&compress]{natbib}
\usepackage{orcidlink}
\usepackage{footmisc}
\usepackage{enumitem}
\usepackage{etoolbox}
\usepackage{calc}
\usepackage{lastpage}
\usepackage{ifthen}

% ── Font ──
""" + font_setup + r"""

% ── Colours ──
\definecolor{SAARJbrown}{HTML}{""" + accent_hex + r"""}
\definecolor{SAARJblue}{HTML}{0070C0}
\definecolor{SAARJgray}{HTML}{CFCDCD}
\definecolor{SAARJdarkgray}{HTML}{3B3838}

% ── Hyperlinks ──
\hypersetup{colorlinks=true,urlcolor=SAARJblue,linkcolor=black,citecolor=black,pdfencoding=auto,unicode=true}
% URL'leri bölme — sığmazsa bütün olarak alt satıra geç
\renewcommand{\UrlBreaks}{}
\renewcommand{\UrlBigBreaks}{}

% ── Geometry ──
\geometry{a4paper,top=1.5cm,bottom=1.5cm,left=1.5cm,right=1.5cm,headheight=1.2cm,headsep=0.4cm,footskip=0.8cm}

% ── Paragraph format ──
\setlength{\parindent}{0pt}
\setlength{\parskip}{4pt}
\renewcommand{\baselinestretch}{1.0}

% ── Satır kırma / Line breaking ──
% Uzun kelimeler ve URL'lerin sayfa kenarına taşmasını önler
\setlength{\emergencystretch}{3em}
\tolerance=800
\hyphenpenalty=50
\exhyphenpenalty=50

% ── Section headings ──
\titleformat{\section}[block]{\fontsize{11}{13}\selectfont\bfseries\centering}{}{0em}{}
\titlespacing*{\section}{0pt}{9pt}{5pt}
\titleformat*{\section}{\fontsize{11}{13}\selectfont\bfseries\centering}
\titleformat{\subsection}[block]{\fontsize{11}{13}\selectfont\bfseries}{}{0em}{}
\titlespacing*{\subsection}{0pt}{7pt}{3pt}
\titleformat{\subsubsection}[block]{\fontsize{11}{13}\selectfont\bfseries\itshape}{}{0em}{}
\titlespacing*{\subsubsection}{0pt}{6pt}{3pt}

% ── Headers & footers ──
\pagestyle{fancy}
\fancyhf{}
\fancyhead[C]{%
  \fontsize{8.5}{10.5}\selectfont
  """ + author_short + r"""\ (""" + year + r""").
  """ + head_title + r""".
  \textit{""" + escape(jname_en) + r"""},
  \textit{""" + volume + r"""}(""" + issue + r"""),
  """ + start_page + r"""--""" + end_page + r"""%
}
\fancyfoot[C]{\fontsize{9}{11}\selectfont\thepage}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}
\fancypagestyle{firstpage}{\fancyhf{}\renewcommand{\headrulewidth}{0pt}}

% ── Caption format ──
\captionsetup{font={small,bf},labelsep=period,justification=centering,skip=4pt}
\captionsetup[table]{position=top}

% ── Footnote format ──
\renewcommand{\footnoterule}{\kern-3pt\hrule width 2cm height 0.4pt\kern 2.6pt}
\setlength{\footnotesep}{4pt}

% ── List format ──
\setlist{leftmargin=1.2cm,itemsep=0pt,parsep=0pt,topsep=2pt,partopsep=0pt}

% ── Bibliography ──
\bibliographystyle{unsrtnat}
\setlength{\bibhang}{1.5em}
\setlength{\bibsep}{3pt}

% ── Article metadata ──
\newcommand{\SAARJyear}{""" + year + r"""}
\newcommand{\SAARJvolume}{""" + volume + r"""}
\newcommand{\SAARJissue}{""" + issue + r"""}
\newcommand{\SAARJstartpage}{""" + start_page + r"""}
\newcommand{\SAARJendpage}{""" + end_page + r"""}
\newcommand{\SAARJDOI}{""" + doi_line + r"""}
\newcommand{\SAARJarticletype}{""" + art_type + r"""}
\newcommand{\SAARJreceived}{""" + received + r"""}
\newcommand{\SAARJaccepted}{""" + accepted + r"""}
\newcommand{\SAARJpublished}{""" + published + r"""}
\newcommand{\SAARJturkishtitle}{""" + tr_title + r"""}
\newcommand{\SAARJturkishabstract}{""" + tr_abs + r"""}
\newcommand{\SAARJturkishkeywords}{""" + tr_kw + r"""}
\newcommand{\SAARJenglishtitle}{""" + en_title + r"""}
\newcommand{\SAARJenglishabstract}{""" + en_abs + r"""}
\newcommand{\SAARJenglishkeywords}{""" + en_kw + r"""}
\newcommand{\SAARJauthorshort}{""" + author_short + r"""}
\newcommand{\SAARJheadtitle}{""" + head_title + r"""}
\newcommand{\SAARJaffiliations}{""" + affiliations + r"""}
\newcommand{\SAARJcorrespondinginfo}{""" + corresponding + r"""}
\newcommand{\SAARJapacitation}{""" + apa_citation + r"""}
\newcommand{\SAARJethicsstatement}{""" + ethics + r"""}

% ── DOI yardımcı makrosu (doi: önekini tekrar etmeden hyperlink) ──
\newcommand{\SAARJdoilink}{\href{https://doi.org/\SAARJDOI}{\SAARJDOI}}

% ── Internal helpers ──
\newcommand{\infolabel}[1]{{\fontsize{9}{11}\selectfont\scshape #1}}
\newcommand{\infosubheading}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\infovalue}[1]{{\fontsize{7.5}{9}\selectfont #1}}
\newcommand{\infoboldlabel}[1]{{\fontsize{7.5}{9}\selectfont\bfseries #1}}
\newcommand{\inforule}{\par\vspace{0pt}\noindent\rule{\linewidth}{0.4pt}\par\vspace{0pt}}

% ── First page command ──
\newcommand{\SAARJfirstpage}[1]{%
  \newgeometry{a4paper,top=""" + f"{top_margin_cm}" + r"""cm,bottom=1.5cm,left=1.5cm,right=1.5cm,headheight=0pt,headsep=0pt,footskip=0.8cm}%
  \thispagestyle{firstpage}%

  \noindent%
  {\setlength{\tabcolsep}{0pt}\setlength{\extrarowheight}{0pt}%
  \begin{tabular}{@{} m{11cm} @{} >{\raggedleft\arraybackslash}m{7cm} @{}}%
    \includegraphics[height=""" + f"{logo_height_cm}" + r"""cm]{""" + logo_stem + r"""}%
    &%
    \raggedleft%
    {\fontsize{8.5}{11}\selectfont\color{SAARJdarkgray}%
      """ + issn_display + r"""
      \href{""" + j_url + r"""}{\color{SAARJblue}\itshape """ + j_url + r"""}""" + (
        r"""\\[2pt]%
      \ifthenelse{\equal{\SAARJDOI}{}}{\textbf{DOI:} \textit{""" + (r'(to be assigned)' if english_only else r'(atanacak\,/\,to be assigned)') + r"""}}{\textbf{DOI:} \SAARJdoilink}"""
        if doi_position == 'top' else ''
      ) + r"""\par}%
  \end{tabular}}%

  \vspace{-1.5mm}%
  \noindent\rule{\textwidth}{0.4pt}%

  {\noindent\setlength{\fboxsep}{1.5pt}\setlength{\fboxrule}{0pt}%
  \colorbox{SAARJgray}{%
    \begin{minipage}{\dimexpr\textwidth-3pt\relax}%
      """ + _build_meta_strip(english_only) + r"""%
    \end{minipage}%
  }}%

  \vspace{1.5mm}%

  {\noindent\centering{\fontsize{10.5}{12}\selectfont\bfseries\itshape \SAARJarticletype}\par}%

  \vspace{1.5mm}%

  """ + _build_titles(english_only) + r"""

  \vspace{1.5mm}%

  {\noindent{\fontsize{10}{12}\selectfont #1}\par}%

  \vspace{0.5mm}%
  \noindent\rule{\textwidth}{0.4pt}%

  """ + _build_abstract_block(english_only, has_tr_abs, has_en_abs, editor_raw) + r"""

  \noindent\rule{\textwidth}{0.4pt}%

  \vspace*{\fill}%
  \noindent\begin{minipage}{\textwidth}%
  \begingroup
    \sloppy\emergencystretch=3em%
    \setlength{\parskip}{1pt}\setlength{\parindent}{0pt}%
    \fontsize{7}{8.5}\selectfont%
    \noindent\rule{\textwidth}{0.4pt}\par%
    \noindent\SAARJaffiliations\par%
    \noindent\rule{\textwidth}{0.2pt}\par%
    \noindent\textit{""" + _build_corresponding_label(english_only, corr_marker) + r"""}\par%
    \noindent\textbf{""" + (r'Suggested Citation:' if english_only else r'Önerilen Atıf / Suggested Citation:') + r"""} \SAARJapacitation\par%
    \noindent{\setlength{\parskip}{0pt}\textbf{""" + (r'Ethics Statement:' if english_only else r'Etik Beyan / Ethics Statement:') + r"""} \SAARJethicsstatement\par}%
  \endgroup%
  \end{minipage}

  % --- Alt bilgi şeridi: sol CC-BY (+ sağ DOI, eğer doi_position == 'bottom') ---
  \vspace{4pt}%
  \noindent\begin{minipage}{\textwidth}%
    \noindent%
    \begin{minipage}[c]{0.50\textwidth}\raggedright%
      \includegraphics[height=0.85cm]{""" + cc_logo_stem + r"""}%
    \end{minipage}%
    \begin{minipage}[c]{0.50\textwidth}\raggedleft%
      """ + (
        r"""{\fontsize{9}{11}\selectfont\color{SAARJdarkgray}%
        \ifthenelse{\equal{\SAARJDOI}{}}{\textbf{DOI:} \textit{""" + (r'(to be assigned)' if english_only else r'(atanacak\,/\,to be assigned)') + r"""}}{\textbf{DOI:} \SAARJdoilink}%
      }"""
        if doi_position == 'bottom' else ''
      ) + r"""%
    \end{minipage}%
  \end{minipage}

  \restoregeometry%
}% end \SAARJfirstpage


% ============================================================
\begin{document}

\SAARJfirstpage{%
  """ + author_block + r"""%
}
% Kapak sayfası sayfa 1'dir ancak numara gösterilmez.
% İkinci sayfa sayfa 2 olarak başlar.

""" + body_text + '\n\n' + refs_tex + (('\n\n' + extended_abstract_tex) if extended_abstract_tex else '') + r"""

\end{document}
"""
    return tex


def build_zip_form(tex_content: str, logo_src: str, figure_file_bytes: dict,
                   journal_settings: dict = None,
                   ccby_src: str = None,
                   ccby_upload: tuple = None) -> bytes:
    """
    Build Overleaf-ready ZIP including:
    - main.tex
    - journal_logo.<ext> (logo)
    - ccby.<ext> (CC-BY 4.0 logo, default or uploaded)
    - figure files (fig_N.ext)
    - README_Overleaf.txt
    """
    js = journal_settings or {}
    logo_stem = js.get('logo_stem', 'journal_logo')
    logo_fn   = js.get('logo_filename', '')   # e.g. "SAARJ_logo.png"

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('main.tex', tex_content.encode('utf-8'))

        if os.path.exists(logo_src):
            # Determine extension of the actual logo file
            ext = logo_src.rsplit('.', 1)[-1].lower() if '.' in logo_src else 'png'
            zf.write(logo_src, logo_stem + '.' + ext)

        # CC-BY logo — upload varsa onu, yoksa varsayılanı ekle
        if ccby_upload and ccby_upload[1]:
            zf.writestr(ccby_upload[0], ccby_upload[1])
        elif ccby_src and os.path.exists(ccby_src):
            with open(ccby_src, 'rb') as _f:
                zf.writestr('ccby.png', _f.read())

        for fkey, (zipname, filebytes) in figure_file_bytes.items():
            zf.writestr(zipname, filebytes)

        readme = (
            "SAARJ LaTeX Formatter — Overleaf Yükleme Rehberi\n"
            "=================================================\n\n"
            "1. Bu ZIP dosyasını açın.\n"
            "2. Overleaf.com → New Project → Upload Project → ZIP'i seçin.\n"
            "3. Menu → Compiler → XeLaTeX seçin.\n"
            "4. Recompile → PDF hazır.\n"
        )
        zf.writestr('README_Overleaf.txt', readme.encode('utf-8'))

    buf.seek(0)
    return buf.read()




